#!/usr/bin/env python3
"""
📚 DOCUMENT TO JSON BATCH PROCESSOR
===================================
Memory-efficient extraction from PDF, DOC, DOCX files in batches.

Supports:
  - PDF files (via pdfplumber)
  - DOCX files (via python-docx)
  - DOC files (via python-docx or conversion)

Features:
  - Processes one page/section at a time
  - Batches content automatically
  - Detects lessons, chapters, sections
  - Extracts text, images, tables
  - Saves incrementally (no memory overload)
  - Works with large documents (100+ pages)

Installation:
  pip install pdfplumber python-docx PyPDF2
"""

import json
import sys
from pathlib import Path
from typing import Dict, List, Generator, Optional
import re


def install_dependencies():
    """Check and install required packages"""
    print("📦 Checking dependencies...")

    required = {
        'pdfplumber': 'For PDF processing',
        'python_docx': 'For DOCX processing',
        'PyPDF2': 'For PDF handling'
    }

    missing = []
    for package, description in required.items():
        try:
            __import__(package)
            print(f"  ✅ {package:20} - {description}")
        except ImportError:
            print(f"  ❌ {package:20} - {description}")
            missing.append(package)

    if missing:
        print(f"\n⚠️  Installing missing packages: {', '.join(missing)}")
        import subprocess
        for package in missing:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q', package])
        print("✅ Dependencies installed!")
    else:
        print("\n✅ All dependencies installed!")


class DocumentBatchProcessor:
    """Process large documents in memory-efficient batches"""

    def __init__(self, document_file: str, batch_size: int = 10, output_dir: str = None):
        self.doc_file = Path(document_file)
        self.batch_size = batch_size
        self.batch_num = 0
        self.total_pages = 0
        self.file_type = self.doc_file.suffix.lower()

        # Validate input
        if not self.doc_file.exists():
            print(f"❌ Error: File not found: {document_file}")
            sys.exit(1)

        if self.file_type not in ['.pdf', '.docx', '.doc']:
            print(f"❌ Error: Unsupported file type: {self.file_type}")
            print("Supported: .pdf, .docx, .doc")
            sys.exit(1)

        # Set output directory
        if output_dir is None:
            output_dir = f"{self.doc_file.stem}_extracted"

        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

        # Statistics
        self.stats = {
            'batches': 0,
            'sections_found': 0,
            'lessons_found': 0,
            'tables_found': 0,
            'total_chars': 0,
            'images_found': 0
        }

    def process_pdf(self) -> Generator[List[Dict], None, None]:
        """Process PDF file in batches"""
        try:
            import pdfplumber
        except ImportError:
            print("❌ pdfplumber not available. Install with: pip install pdfplumber")
            return

        try:
            with pdfplumber.open(self.doc_file) as pdf:
                self.total_pages = len(pdf.pages)
                current_batch = []

                for page_num, page in enumerate(pdf.pages, 1):
                    section = self._extract_pdf_page(page, page_num)
                    current_batch.append(section)

                    if len(current_batch) >= self.batch_size:
                        yield current_batch
                        current_batch = []

                if current_batch:
                    yield current_batch

        except Exception as e:
            print(f"❌ Error processing PDF: {e}")

    def process_docx(self) -> Generator[List[Dict], None, None]:
        """Process DOCX file in batches"""
        try:
            from docx import Document
        except ImportError:
            print("❌ python-docx not available. Install with: pip install python-docx")
            return

        try:
            doc = Document(self.doc_file)
            current_batch = []
            section_num = 0

            for para_idx, paragraph in enumerate(doc.paragraphs, 1):
                # Detect section headers
                is_header = self._is_section_header(paragraph.text)

                section = {
                    'index': para_idx,
                    'type': 'heading' if is_header else 'paragraph',
                    'text': paragraph.text,
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'char_count': len(paragraph.text)
                }

                if is_header:
                    section_num += 1
                    section['section_number'] = section_num
                    self.stats['sections_found'] += 1

                current_batch.append(section)
                self.stats['total_chars'] += len(paragraph.text)

                if len(current_batch) >= self.batch_size:
                    yield current_batch
                    current_batch = []

            if current_batch:
                yield current_batch

            self.total_pages = len(doc.paragraphs)

        except Exception as e:
            print(f"❌ Error processing DOCX: {e}")

    def _extract_pdf_page(self, page, page_num: int) -> Dict:
        """Extract content from PDF page"""
        text = page.extract_text() or ""
        tables = page.extract_tables() or []

        section = {
            'page': page_num,
            'type': 'pdf_page',
            'text': text,
            'text_length': len(text),
            'tables_count': len(tables),
            'has_images': len(page.images) > 0,
            'lessons': self._detect_lessons(text, page_num)
        }

        self.stats['total_chars'] += len(text)
        self.stats['tables_found'] += len(tables)
        self.stats['lessons_found'] += len(section['lessons'])

        return section

    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header"""
        if not text or len(text) < 3:
            return False

        # Check style indicators
        headers = [
            text.isupper() and len(text) < 100,
            text.startswith('Chapter'),
            text.startswith('Lecția'),
            text.startswith('Unit'),
            text.startswith('Unitatea'),
            re.match(r'^(Lecția|Chapter|Unit|Unitatea|Capitolul)\s+\d+', text, re.IGNORECASE)
        ]

        return any(headers)

    def _detect_lessons(self, text: str, location: int) -> List[Dict]:
        """Detect lessons in text"""
        lessons = []
        patterns = [
            r'^Lecția\s+(\d+)[:\s]+(.+?)$',
            r'^Chapter\s+(\d+)[:\s]+(.+?)$',
            r'^Unit\s+(\d+)[:\s]+(.+?)$',
            r'^Unitatea\s+(\d+)[:\s]+(.+?)$',
        ]

        for line in text.split('\n'):
            for pattern in patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    lessons.append({
                        'number': int(match.group(1)),
                        'title': match.group(2).strip(),
                        'location': location
                    })

        return lessons

    def process_batch(self, sections: List[Dict]) -> Dict:
        """Process a batch of sections"""
        return {
            'batch_number': self.batch_num,
            'section_count': len(sections),
            'sections': sections,
            'summary': {
                'total_chars': sum(s.get('text_length', len(s.get('text', ''))) for s in sections),
                'total_lessons': sum(len(s.get('lessons', [])) for s in sections),
                'location_range': f"{sections[0].get('index', sections[0].get('page'))}-{sections[-1].get('index', sections[-1].get('page'))}"
            }
        }

    def save_batch(self, batch_data: Dict) -> Path:
        """Save batch to JSON"""
        batch_file = self.output_dir / f"batch_{self.batch_num:03d}.json"

        try:
            with open(batch_file, 'w', encoding='utf-8') as f:
                json.dump(batch_data, f, ensure_ascii=False, indent=2)
            return batch_file
        except IOError as e:
            print(f"❌ Error saving batch: {e}")
            return None

    def print_summary(self):
        """Print processing summary"""
        print("\n" + "="*80)
        print("📊 DOCUMENT EXTRACTION COMPLETE")
        print("="*80)
        print(f"✅ Batches created:      {self.stats['batches']}")
        print(f"📄 Sections processed:   {self.stats['sections_found']}")
        print(f"📝 Lessons detected:     {self.stats['lessons_found']}")
        print(f"📊 Total characters:     {self.stats['total_chars']:,}")
        print(f"📋 Tables found:         {self.stats['tables_found']}")
        print(f"📂 Output directory:     {self.output_dir.absolute()}")
        print("="*80 + "\n")

    def process_all(self) -> bool:
        """Main processing loop"""
        print("\n" + "="*80)
        print("🚀 DOCUMENT TO JSON BATCH PROCESSOR")
        print("="*80)
        print(f"📖 File:             {self.doc_file}")
        print(f"📄 Type:             {self.file_type}")
        print(f"📦 Batch size:       {self.batch_size} sections")
        print(f"📂 Output dir:       {self.output_dir}")
        print("="*80 + "\n")

        # Choose processor based on file type
        if self.file_type == '.pdf':
            batches = self.process_pdf()
        elif self.file_type in ['.docx', '.doc']:
            batches = self.process_docx()
        else:
            print(f"❌ Unsupported file type: {self.file_type}")
            return False

        try:
            for batch in batches:
                batch_data = self.process_batch(batch)
                batch_file = self.save_batch(batch_data)

                if batch_file:
                    print(f"✅ Batch {self.batch_num:02d}: {len(batch):3d} sections → {batch_file.name}")
                    self.stats['batches'] += 1
                    self.batch_num += 1

            self.print_summary()
            return True

        except Exception as e:
            print(f"❌ Error: {e}")
            return False


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description='Extract documents (PDF/DOCX) to JSON in batches',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process PDF with default batch size
  %(prog)s manual.pdf
  
  # Process DOCX with custom batch size
  %(prog)s document.docx -b 20
  
  # Custom output directory
  %(prog)s file.pdf -o output_json
  
  # Install dependencies first
  %(prog)s --install-deps
        """
    )

    parser.add_argument('document', nargs='?', help='Document file (PDF, DOCX, DOC)')
    parser.add_argument('-b', '--batch-size', type=int, default=10,
                        help='Sections per batch (default: 10)')
    parser.add_argument('-o', '--output', help='Output directory')
    parser.add_argument('--install-deps', action='store_true',
                        help='Install required dependencies')

    args = parser.parse_args()

    if args.install_deps:
        install_dependencies()
        return

    if not args.document:
        parser.print_help()
        return

    processor = DocumentBatchProcessor(
        document_file=args.document,
        batch_size=args.batch_size,
        output_dir=args.output
    )

    success = processor.process_all()
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()

