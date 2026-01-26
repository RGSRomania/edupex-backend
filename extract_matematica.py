#!/usr/bin/env python3
"""
Extract lesson content from Manual MATE.doc (Math manual for Clasa a V-a)
This script reads the .doc file and extracts structured lesson content
including summaries, examples, notes, and important concepts.
"""

from docx import Document
import json
import re

def extract_matematica_lessons():
    """
    Extract lessons from Manual MATE.doc
    Returns a list of lessons with their content
    """

    doc_path = "/Users/mdica/PycharmProjects/EduPex/Manuale word/Clasa a V-A/Manual MATE.doc"

    try:
        doc = Document(doc_path)
        print(f"✅ Successfully opened: {doc_path}")
        print(f"✅ Document has {len(doc.paragraphs)} paragraphs")
        print(f"✅ Document has {len(doc.tables)} tables")
        print("\n" + "="*60)

        # Print first 50 paragraphs to understand structure
        print("\n📄 First 50 paragraphs of the document:")
        print("-"*60)

        for i, para in enumerate(doc.paragraphs[:100]):
            text = para.text.strip()
            if text:  # Only print non-empty paragraphs
                # Show formatting style
                style = para.style.name if para.style else "None"
                level = para.paragraph_format.outline_level if para.paragraph_format else 0

                if i % 10 == 0 or len(text) > 100:
                    print(f"\n[Para {i}] Level:{level} Style:{style}")
                    print(f"  Text: {text[:100]}...")
                else:
                    if len(text) < 80:
                        print(f"[{i}] {text}")

        print("\n" + "="*60)
        print("\n📊 Document Structure Analysis:")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")

        # Count by style
        styles = {}
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "None"
            styles[style_name] = styles.get(style_name, 0) + 1

        print("\nParagraph styles used:")
        for style, count in sorted(styles.items(), key=lambda x: x[1], reverse=True):
            print(f"  {style}: {count}")

        return doc

    except Exception as e:
        print(f"❌ Error: {e}")
        return None


if __name__ == '__main__':
    doc = extract_matematica_lessons()

