#!/usr/bin/env python3
"""
Extract clean, structured content from Word documents (.doc/.docx)
This produces MUCH cleaner text than PDF extraction.
"""

import os
import sys
from pathlib import Path

# Try to import python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. Installing...")
    os.system("pip install python-docx")
    from docx import Document
    DOCX_AVAILABLE = True

# For .doc files (older format), try python-docx2docx or direct extraction
try:
    from docx2docx import convert
    DOC_CONVERTER = True
except ImportError:
    DOC_CONVERTER = False

def extract_from_docx(file_path):
    """Extract clean text from .docx files"""
    print(f"\n📖 Extracting from DOCX: {file_path}")

    try:
        doc = Document(file_path)

        sections = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Detect section headers (usually bold or larger)
            if para.style.name.startswith('Heading'):
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'title': text,
                    'content': [],
                    'level': int(para.style.name.replace('Heading', ''))
                }
            elif current_section:
                current_section['content'].append(text)
            else:
                if not sections:
                    current_section = {
                        'title': 'Introduction',
                        'content': [text],
                        'level': 0
                    }
                else:
                    current_section['content'].append(text)

        if current_section:
            sections.append(current_section)

        return sections

    except Exception as e:
        print(f"❌ Error extracting DOCX: {e}")
        return None

def convert_doc_to_docx(doc_path):
    """Convert .doc to .docx for easier processing"""
    print(f"\n🔄 Converting .doc to .docx: {doc_path}")

    docx_path = doc_path.replace('.doc', '.docx')

    # Try using MS Word or LibreOffice
    import subprocess

    # Try using 'libreoffice' command line
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir',
             os.path.dirname(doc_path), doc_path],
            capture_output=True,
            timeout=60
        )

        if result.returncode == 0 and os.path.exists(docx_path):
            print(f"✅ Converted to: {docx_path}")
            return docx_path
    except Exception as e:
        print(f"⚠️  LibreOffice conversion failed: {e}")

    # Fallback: try unoconv
    try:
        result = subprocess.run(
            ['unoconv', '-f', 'docx', doc_path],
            capture_output=True,
            timeout=60
        )
        if result.returncode == 0 and os.path.exists(docx_path):
            print(f"✅ Converted using unoconv: {docx_path}")
            return docx_path
    except Exception as e:
        print(f"⚠️  Unoconv not available: {e}")

    return None

def extract_from_doc(file_path):
    """Extract from .doc file (try multiple methods)"""
    print(f"\n📖 Processing .doc file: {file_path}")

    # First, try to convert to docx
    docx_path = convert_doc_to_docx(file_path)

    if docx_path and os.path.exists(docx_path):
        return extract_from_docx(docx_path)

    # Fallback: try direct extraction from .doc
    try:
        from docx import Document
        doc = Document(file_path)
        sections = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append({
                    'content': [text],
                    'level': 0
                })

        return sections
    except Exception as e:
        print(f"❌ Could not extract from .doc: {e}")
        print("   Try installing: pip install python-docx2docx")
        return None

def format_output(sections, output_format='text'):
    """Format extracted sections into readable output"""

    if output_format == 'text':
        output = []
        for section in sections:
            title = section.get('title', 'Section')
            level = section.get('level', 0)
            indent = "  " * level

            output.append(f"\n{indent}{'#' * (level + 1)} {title}")
            output.append("=" * 60)

            content = section.get('content', [])
            for para in content:
                output.append(f"{indent}{para}")

        return "\n".join(output)

    elif output_format == 'json':
        import json
        return json.dumps(sections, indent=2, ensure_ascii=False)

    return sections

def extract_document(file_path, output_format='text', output_file=None):
    """Main function to extract from any document"""

    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return None

    file_ext = Path(file_path).suffix.lower()

    print(f"\n{'='*70}")
    print(f"📄 DOCUMENT EXTRACTION TOOL")
    print(f"{'='*70}")
    print(f"File: {file_path}")
    print(f"Format: {file_ext}")
    print(f"Output: {output_format}")

    # Extract based on file type
    if file_ext == '.docx':
        sections = extract_from_docx(file_path)
    elif file_ext == '.doc':
        sections = extract_from_doc(file_path)
    else:
        print(f"❌ Unsupported format: {file_ext}")
        return None

    if not sections:
        print("❌ Extraction failed!")
        return None

    # Format output
    formatted = format_output(sections, output_format)

    # Print summary
    print(f"\n✅ Extraction successful!")
    print(f"   Found {len(sections)} sections")
    print(f"\n{'='*70}")
    print("PREVIEW (first 2000 characters):")
    print(f"{'='*70}")

    if isinstance(formatted, str):
        print(formatted[:2000])
    else:
        print(str(formatted)[:2000])

    # Save to file if specified
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            if isinstance(formatted, str):
                f.write(formatted)
            else:
                import json
                f.write(json.dumps(formatted, indent=2, ensure_ascii=False))
        print(f"\n💾 Saved to: {output_file}")

    return formatted

if __name__ == '__main__':
    # Example usage
    if len(sys.argv) < 2:
        # Default: extract Matematica manual
        file_path = '/Users/mdica/PycharmProjects/EduPex/Manuale word/Clasa a V-A/Manual MATE.doc'
        output_file = '/Users/mdica/PycharmProjects/EduPex/EXTRACTED_CLEAN_MANUAL.txt'
    else:
        file_path = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else None

    extract_document(file_path, output_format='text', output_file=output_file)

